from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from .structs import Table, NumberedList, BulletedList, Image


class DOCX:
    def __init__(
        self
    ):
        """
        Initializes a DOCX object
        Инициализирует объект DOCX
        """
        self.__docx = Document()
        self.__font_name = "Arial"
        self.__font_size = 14
        self.__alignment = "left"
        self.__font_styles = {
            "bold": False,
            "italic": False,
            "underline": False
        }
        self.set_settings(self.__font_name, self.__font_size, self.__alignment, **self.__font_styles)
        self.__line_spacing = 1
        self.__paragraph_spacing = 1
        self.__first_addition = True
        self.__fixed_paragraph_spacing = None
        self.__increment = 1

    def set_font_name(
        self,
        font_name: str
    ) -> None:
        """
        Sets the font
        Устанавливает шрифт

        Args:
            font_name (str):
                The name of the font to be set
                Имя шрифта, которое нужно установить
        """
        self.__font_name = font_name

    def set_font_size(
        self,
        font_size: float
    ) -> None:
        """
        Sets the font size
        Устанавливает размер шрифта

        Args:
            font_size (float):
                The size of the font to be set
                Размер шрифта, который нужно установить
        """
        self.__font_size = font_size

    def set_alignment(
        self,
        alignment: str
    ) -> None:
        """
        Sets the alignment of the text
        Устанавливает выравнивание текста

        Args:
            alignment (str):
                The alignment of the text. Valid values are 'left', 'right', 'center', 'justify'
                Выравнивание текста. Допустимые значения: 'left', 'right', 'center', 'justify'
        """
        self.__alignment = alignment

    def set_font_style(
        self,
        bold: bool = None,
        italic: bool = None,
        underline: bool = None
    ) -> None:
        """
        Sets the font style
        Устанавливает стиль шрифта

        Args:
            bold (bool, optional):
                Whether the text should be bold. If not specified, the current bold setting remains unchanged
                Определяет, должен ли текст быть полужирным. Если не указано, текущая настройка полужирного остается без изменений
            italic (bool, optional):
                Whether the text should be italic. If not specified, the current italic setting remains unchanged
                Определяет, должен ли текст быть курсивом. Если не указано, текущая настройка курсива остается без изменений
            underline (bool, optional):
                Whether the text should be underlined. If not specified, the current underline setting remains unchanged
                Определяет, должен ли текст быть подчеркнутым. Если не указано, текущая настройка подчеркивания остается без изменений
        """
        if bold is not None:
            self.__font_styles["bold"] = bold
        if italic is not None:
            self.__font_styles["italic"] = italic
        if underline is not None:
            self.__font_styles["underline"] = underline

    def set_settings(
        self,
        font_name: str = None,
        font_size: float = None,
        alignment: str = None,
        bold: bool = None,
        italic: bool = None,
        underline: bool = None
    ) -> None:
        """
        Sets the settings for document
        Устанавливает настройки для документа

        Args:
            font_name (str, optional):
                The name of the font to be set. If not specified, the current font name remains unchanged
                Имя шрифта, которое нужно установить. Если не указано, текущее имя шрифта остается без изменений
            font_size (float, optional):
                The size of the font to be set. If not specified, the current font size remains unchanged
                Размер шрифта, который нужно установить. Если не указано, текущий размер шрифта остается без изменений
            alignment (str, optional):
                The alignment of the text. Valid values are 'left', 'right', 'center', 'justify'. If not specified, the current alignment remains unchanged
                Выравнивание текста. Допустимые значения: 'left', 'right', 'center', 'justify'. Если не указано, текущее выравнивание остается без изменений
            bold (bool, optional):
                Whether the text should be bold. If not specified, the current bold setting remains unchanged
                Определяет, должен ли текст быть полужирным. Если не указано, текущая настройка полужирного остается без изменений
            italic (bool, optional):
                Whether the text should be italicized. If not specified, the current italic setting remains unchanged
                Определяет, должен ли текст быть курсивом. Если не указано, текущая настройка курсива остается без изменений
            underline (bool, optional):
                Whether the text should be underlined. If not specified, the current underline setting remains unchanged
                Определяет, должен ли текст быть подчеркнутым. Если не указано, текущая настройка подчеркивания остается без изменений
        """
        if font_name is not None:
            self.set_font_name(font_name)
        if font_size is not None:
            self.set_font_size(font_size)
        if alignment is not None:
            self.set_alignment(alignment)
        self.set_font_style(bold, italic, underline)

    def set_margin_top(
        self,
        top: float
    ) -> None:
        """
        Sets the top margin of document
        Устанавливает верхний отступ документа

        Args:
            top (float):
                The value of the top margin to be set
                Значение верхнего поля, которое нужно установить
        """
        for section in self.__docx.sections:
            section.top_margin = Cm(top)

    def set_margin_right(
        self,
        right: float
    ) -> None:
        """
        Sets the right margin of document
        Устанавливает правый отступ документа

        Args:
            right (float):
                The value of the right margin to be set
                Значение правого поля, которое нужно установить
        """
        for section in self.__docx.sections:
            section.right_margin = Cm(right)

    def set_margin_bottom(
        self,
        bottom: float
    ) -> None:
        """
        Sets the bottom margin of document
        Устанавливает нижний отступ документа

        Args:
            bottom (float):
                The value of the bottom margin to be set
                Значение нижнего поля, которое нужно установить
        """
        for section in self.__docx.sections:
            section.bottom_margin = Cm(bottom)

    def set_margin_left(
        self,
        left: float
    ) -> None:
        """
        Sets the left margin of document
        Устанавливает левый отступ документа

        Args:
            left (float):
                The value of the left margin to be set
                Значение левого поля, которое нужно установить
        """
        for section in self.__docx.sections:
            section.left_margin = Cm(left)

    def set_margin(
        self,
        top: float = None,
        right: float = None,
        bottom: float = None,
        left: float = None
    ) -> None:
        """
        Sets the margins of document
        Устанавливает отступы документа

        Args:
            top (float, optional):
                The value of the top margin to be set. If not specified, the top margin remains unchanged
                Значение верхнего поля, которое нужно установить. Если не указано, верхнее поле остается без изменений
            right (float, optional):
                The value of the right margin to be set. If not specified, the right margin remains unchanged
                Значение правого поля, которое нужно установить. Если не указано, правое поле остается без изменений
            bottom (float, optional):
                The value of the bottom margin to be set. If not specified, the bottom margin remains unchanged
                Значение нижнего поля, которое нужно установить. Если не указано, нижнее поле остается без изменений
            left (float, optional):
                The value of the left margin to be set. If not specified, the left margin remains unchanged
                Значение левого поля, которое нужно установить. Если не указано, левое поле остается без изменений
        """
        if top is not None:
            self.set_margin_top(top)
        if right is not None:
            self.set_margin_right(right)
        if bottom is not None:
            self.set_margin_bottom(bottom)
        if left is not None:
            self.set_margin_left(left)

    def set_line_spacing(
        self,
        spacing_k: float = 1
    ) -> None:
        """
        Sets the line spacing of document
        Устанавливает межстрочный интервал документа

        Args:
            spacing_k (float):
                The line spacing factor to be set. Single interval by default
                Фактор межстрочного интервала, который нужно установить. По умолчанию одинарный интервал
        """
        self.__line_spacing = spacing_k

    def set_paragraph_spacing(
        self,
        spacing_k: float = 1
    ) -> None:
        """
        Sets the paragraph spacing of document
        Устанавливает интервал между абзацами документа

        Args:
            spacing_k (float):
                The paragraph spacing factor to be set. Single interval by default
                Фактор интервала между абзацами, который нужно установить. По умолчанию одинарный интервал
        """
        self.__paragraph_spacing = spacing_k

    def set_fixed_paragraph_spacing(
        self,
        spacing: float = None
    ) -> None:
        """
        Sets the fixed paragraph spacing of document
        Фиксирует интервал между абзацами документа

        Args:
            spacing (float, optional):
                The fixed spacing value to be set. If not specified, removes the fixed interval
                Значение фиксированного интервала, которое нужно установить. Если не указано, убирает фиксированный интервал
        """
        if spacing is not None:
            self.__fixed_paragraph_spacing = spacing
        else:
            self.__fixed_paragraph_spacing = None

    def add_paragraph(
        self,
        text: str,
        spacing: float = None,
        font_name: str = None,
        font_size: float = None,
        alignment: str = None,
        bold: bool = None,
        italic: bool = None,
        underline: bool = None
    ) -> None:
        """
        Adds a paragraph to document
        Добавляет параграф в документ

        Args:
            text (str):
                The text content of the paragraph
                Текстовое содержимое абзаца
            spacing (float, optional):
                The spacing to be added before the paragraph. If not specified, default settings are used
                Интервал, добавляемый перед абзацем. Если не указан, используются настройки по умолчанию
            font_name (str, optional):
                The name of the font to be used. If not specified, default settings are used
                Имя используемого шрифта. Если не указано, используются настройки по умолчанию
            font_size (float, optional):
                The size of the font. If not specified, default settings are used
                Размер шрифта. Если не указано, используются настройки по умолчанию
            alignment (str, optional):
                The alignment of the text. Can be 'left', 'right', 'center', or 'justify'. If not specified, default settings are used
                Выравнивание текста. Может быть 'left', 'right', 'center' или 'justify'. Если не указано, используются настройки по умолчанию
            bold (bool, optional):
                Whether the text should be bold. If not specified, default settings are used
                Жирный ли текст. Если не указано, используются настройки по умолчанию
            italic (bool, optional):
                Whether the text should be italic. If not specified, default settings are used
                Курсивный ли текст. Если не указано, используются настройки по умолчанию
            underline (bool, optional):
                Whether the text should be underlined. If not specified, default settings are used
                Подчеркнут ли текст. Если не указано, используются настройки по умолчанию
        """
        self._add_text(
            text,
            spacing,
            font_name,
            font_size,
            alignment,
            bold,
            italic,
            underline
        )

    def add_numbered_list(
        self,
        numbered_list: NumberedList,
        spacing: float = None
    ) -> None:
        """
        Adds a numbered list to document
        Добавляет нумерованный список в документ

        Args:
            numbered_list (NumberedList):
                The list object to be added
                Объект списка, который нужно добавить
            spacing (float, optional):
                The interval to add before the start of the list. If not specified, the default settings are used
                Интервал, добавляемый перед началом списка. Если не указан, используются настройки по умолчанию
        """
        style_name = f"LIST_{self.__increment}"
        self.__increment += 1
        self.__docx.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        self.__docx.styles[style_name].base_style = self.__docx.styles["List Number"]
        self.__docx.styles[style_name].paragraph_format.number_format = numbered_list.number_string.replace("{i}", "\t")
        self.__docx.styles[style_name].paragraph_format.left_indent = Cm(numbered_list.indent)
        self.__docx.styles[style_name].paragraph_format.first_line_indent = 0
        ns = numbered_list.number_settings
        self.__docx.styles[style_name].font.name = ns.font_name or self.__font_name
        self.__docx.styles[style_name].font.size = Pt(ns.font_size or self.__font_size)
        self.__docx.styles[style_name].paragraph_format.alignment = self._make_alignment(ns.alignment or self.__alignment)
        self.__docx.styles[style_name].font.bold = ns.bold or self.__font_styles["bold"]
        self.__docx.styles[style_name].font.italic = ns.italic or self.__font_styles["italic"]
        self.__docx.styles[style_name].font.underline = ns.underline or self.__font_styles["underline"]
        for i in range(len(numbered_list.data)):
            if i == 0:
                if spacing is not None:
                    up_spacing = spacing
                elif self.__fixed_paragraph_spacing is not None:
                    up_spacing = self.__fixed_paragraph_spacing
                else:
                    up_spacing = self.__paragraph_spacing
            else:
                up_spacing = numbered_list.between_spacing
            self._add_text(
                numbered_list.data[i],
                up_spacing,
                **numbered_list.settings[i].get(),
                style=style_name
            )

    def add_bulleted_list(
        self,
        bulleted_list: BulletedList,
        spacing: float = None
    ) -> None:
        """
        Adds a bulleted list to document
        Добавляет маркированный список в документ

        Args:
            bulleted_list (BulletedList):
                The list object to be added
                Объект списка, который нужно добавить
            spacing (float, optional):
                The interval to add before the start of the list. If not specified, the default settings are used
                Интервал, добавляемый перед началом списка. Если не указан, используются настройки по умолчанию
        """
        style_name = f"LIST_{self.__increment}"
        self.__increment += 1
        self.__docx.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        self.__docx.styles[style_name].base_style = self.__docx.styles["List Bullet"]
        self.__docx.styles[style_name].paragraph_format.left_indent = Cm(bulleted_list.indent)
        self.__docx.styles[style_name].paragraph_format.first_line_indent = 0
        bs = bulleted_list.bullet_settings
        self.__docx.styles[style_name].font.name = bs.font_name or self.__font_name
        self.__docx.styles[style_name].font.size = Pt(bs.font_size or self.__font_size)
        self.__docx.styles[style_name].paragraph_format.alignment = self._make_alignment(bs.alignment or self.__alignment)
        self.__docx.styles[style_name].font.bold = bs.bold or self.__font_styles["bold"]
        self.__docx.styles[style_name].font.italic = bs.italic or self.__font_styles["italic"]
        self.__docx.styles[style_name].font.underline = bs.underline or self.__font_styles["underline"]
        for i in range(len(bulleted_list.data)):
            if i == 0:
                if spacing is not None:
                    up_spacing = spacing
                elif self.__fixed_paragraph_spacing is not None:
                    up_spacing = self.__fixed_paragraph_spacing
                else:
                    up_spacing = self.__paragraph_spacing
            else:
                up_spacing = bulleted_list.between_spacing
            self._add_text(
                bulleted_list.data[i],
                up_spacing,
                **bulleted_list.settings[i].get(),
                style=style_name
            )

    def add_image(
        self,
        image: Image,
        image_width,
        image_height,
        image_alignment: str = None,
        image_spacing: float = None,
        caption: str = None,
        caption_spacing: float = None
    ) -> None:
        """
        Adds an image to document
        Добавляет изображение в документ

        Args:
            image (Image):
                The image object to be added
                Объект изображения, который нужно добавить
            image_width (float):
                The width of the image in document
                Ширина изображения в документе
            image_height (float):
                The height of the image in document
                Высота изображения в документе
            image_alignment (str, optional):
                The alignment of the image. Can be 'left', 'right', 'center', or 'justify'. If not specified, default settings are used
                Выравнивание изображения. Может быть 'left', 'right', 'center' или 'justify'. Если не указано, используются настройки по умолчанию
            image_spacing (float, optional):
                The spacing to be added before the image. If not specified, the default settings are used
                Интервал, добавляемый перед изображением. Если не указан, используются настройки по умолчанию
            caption (str, optional):
                The text to be added below the image. If not specified, the text is not added
                Текст, который нужно добавить под изображением. Если не указан, текст не добавляется
            caption_spacing (float, optional):
                The spacing to be added before the text. If not specified, the default settings are used
                Интервал, добавляемый перед текстом. Если не указан, используются настройки по умолчанию
        """
        settings = image.settings
        img_alignment = self.__alignment if image_alignment is None else image_alignment
        paragraph = self.__docx.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image.image_path, width=Cm(image_width), height=Cm(image_height))
        p_format = paragraph.paragraph_format
        p_font = run.font
        p_format.alignment = self._make_alignment(img_alignment)
        p_format.space_after = 0
        if image_spacing is not None:
            p_format.space_before = Cm(image_spacing)
        elif self.__fixed_paragraph_spacing is not None:
            p_format.space_before = Cm(self.__fixed_paragraph_spacing)
        else:
            p_format.space_before = int(self.__paragraph_spacing * p_font.size)
        if caption is not None:
            self._add_text(caption, caption_spacing, **settings.get())

    def add_table(
        self,
        table: Table,
        table_width: float,
        table_height: float,
        table_alignment: str = None,
        spacing: float = None
    ) -> None:
        """
        Adds a table to document
        Добавляет таблицу в документ

        Args:
            table (Table):
                The table object to be added
                Объект таблицы, который нужно добавить
            table_width (float)
                The width of the table in document
                Ширина таблицы в документе
            table_height (float)
                The height of the table in document
                Высота таблицы в документе
            table_alignment (str, optional)
                The alignment of the table. Can be 'left', 'right', 'center', or 'justify'. If not specified, default settings are used
                Выравнивание таблицы. Может быть 'left', 'right', 'center' или 'justify'. Если не указано, используются настройки по умолчанию
            spacing (float, optional):
                The spacing to be added before the table. If not specified, the default settings are used
                Интервал, добавляемый перед таблицей. Если не указан, используются настройки по умолчанию
        """
        paragraph = self.__docx.add_paragraph()
        run = paragraph.add_run()
        p_format = paragraph.paragraph_format
        p_font = run.font
        p_format.space_after = 0
        if spacing is not None:
            p_format.space_before = Cm(spacing)
        elif self.__fixed_paragraph_spacing is not None:
            p_format.space_before = Cm(self.__fixed_paragraph_spacing)
        else:
            p_format.space_before = int(self.__paragraph_spacing * p_font.size)
        cr_table = self.__docx.add_table(rows=len(table.data), cols=len(table.data[0]))
        cr_table.style = "Table Grid"
        cr_table.alignment = self._make_alignment(table_alignment or self.__alignment)
        for i, row in enumerate(table.data):
            cr_table.rows[i].height = Cm(table_height) // len(table.data)
            for j, cell_text in enumerate(table.data[i]):
                settings = table.settings[i][j]
                cell = cr_table.cell(i, j)
                cell.width = Cm(table_width) // len(table.data[0])
                cell.text = cell_text
                cell_p_format = cell.paragraphs[0]
                cell_p_run = cell_p_format.runs[0]
                cell_p_font = cell_p_run.font
                cell_p_font.name = settings.font_name or self.__font_name
                cell_p_font.size = Pt(settings.font_size or self.__font_size)
                cell_p_format.alignment = self._make_alignment(settings.alignment or self.__alignment)
                cell_p_run.bold = settings.bold or self.__font_styles["bold"]
                cell_p_run.italic = settings.italic or self.__font_styles["italic"]
                cell_p_run.underline = settings.underline or self.__font_styles["underline"]

    def save(
        self,
        file_path: str
    ) -> None:
        """
        Saves document to a file
        Сохраняет документ в файл

        Args:
            file_path (str):
                The path to the file to save
                Путь к файлу для сохранения
        """
        self.__docx.save(file_path)

    def _add_text(
        self,
        text: str,
        spacing: float = None,
        font_name: str = None,
        font_size: float = None,
        alignment: str = None,
        bold: bool = None,
        italic: bool = None,
        underline: bool = None,
        style: str = None
    ) -> None:
        paragraph = self.__docx.add_paragraph(style=style)
        p_format = paragraph.paragraph_format
        p_font = paragraph.add_run(text).font
        p_font.name = font_name or self.__font_name
        p_font.size = Pt(font_size or self.__font_size)
        p_format.alignment = self._make_alignment(alignment or self.__alignment)
        p_font.bold = bold or self.__font_styles["bold"]
        p_font.italic = italic or self.__font_styles["italic"]
        p_font.underline = underline or self.__font_styles["underline"]
        if spacing is not None:
            p_format.space_before = Cm(spacing)
        elif self.__fixed_paragraph_spacing is not None:
            p_format.space_before = Cm(self.__fixed_paragraph_spacing)
        else:
            print(int(self.__paragraph_spacing * p_font.size))
            p_format.space_before = int(self.__paragraph_spacing * p_font.size)
        p_format.space_after = 0

    @staticmethod
    def _make_alignment(
            alignment: str
    ) -> WD_ALIGN_PARAGRAPH:
        if alignment == "left":
            return WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "justify":
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        raise TypeError(f"Alignment '{alignment}' is not valid. Valid values are 'left', 'right', 'center', 'justify'.")
